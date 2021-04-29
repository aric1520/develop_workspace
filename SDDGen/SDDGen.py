import re
import SDDGen_input_UI
import SDDGen_output_UI
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

filename = SDDGen_input_UI.path.get()
savepath = SDDGen_output_UI.path.get()
outputname = SDDGen_output_UI.filename.get()

file_content = ""
with open(filename, 'r') as f:
    a = f.readline()
    while a != '':
        file_content += a.strip()
        a = f.readline()

    modules = []
    results = re.findall(r'函数功能: (.+?)\*.+?@process(.+?)@process.+?\*/(\w+?)\s(\w+)\((.*?)\).*?', file_content)
    for result in results:
        args = []

        arg_list = result[4].split(',')
        for arg in arg_list:
            a = arg.strip().split()
            if len(a) > 1:
                args.append({'type': a[0], 'name': a[1]})

        modules.append({
            'desc': result[0],
            'process': result[1],
            'return': result[2],
            'name': result[3],
            'args': args
        })
    tmps = re.findall(r'/\*(.+?)/\*end-func', file_content)
    funcs = []
    for tmp in tmps:
        func_s = ','
        func_result = re.findall(r'.+?/\*func\*/(.+?)./\*func\*/', tmp)
        funcs.append(func_s.join(func_result))

document = Document()
h1 = document.add_heading("", level=1)
r1 = h1.add_run("1软件详细设计")
r1.font.name = 'Times New Roman'
r1.font.size = Pt(14)
r1.bold = False
r1._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
r1.font.color.rgb = RGBColor(0,0,0)

cnt = 1
for m in modules:
    strcnt = cnt
    h2 = document.add_heading("", level=2)
    r2 = h2.add_run('1.'+ str(strcnt) + m['name'])
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(13)
    r2.bold = False
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    r2.font.color.rgb = RGBColor(0,0,0)

    h3 = document.add_heading("", level=3)
    r3 = h3.add_run('1.'+ str(strcnt) + '.1' + "功能概述")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.bold = False
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    r3.font.color.rgb = RGBColor(0, 0, 0)

    p1 = document.add_paragraph()
    r_p1 = p1.add_run("  " + m['desc'])
    r_p1.font.name = 'Times New Roman'
    r_p1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r_p1.font.size = Pt(12)
    r_p1.font.color.rgb = RGBColor(0, 0, 0)

    h3 = document.add_heading("", level=3)
    r3 = h3.add_run('1.' + str(strcnt) + '.2' + "函数输入")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.bold = False
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    r3.font.color.rgb = RGBColor(0, 0, 0)

    row_num_1 = 1
    for a in m['args']:
        row_num_1 = row_num_1 + 1    # 函数输入表格行数为row_num
    table1 = document.add_table(rows=row_num_1, cols=5, style ='Table Grid')   # 添加函数输入表格
    table1.autofit = True
    table1.cell(0, 0).text = "序号"
    table1.cell(0, 1).text = "数据名称"
    table1.cell(0, 2).text = "数据类型"
    table1.cell(0, 3).text = "单位"
    table1.cell(0, 4).text = "数据范围"
    for x in range(row_num_1 - 1):
        table1.cell(x+1, 0).text = str(x+1)   # 表格每行序号
    x = 1
    for a in m['args']:             # 填入数据名称和数据类型
        table1.cell(x, 1).text = a['name']
        table1.cell(x, 2).text = a['type']
        x = x+1

    for row in table1.rows:  # 设置表格中字体格式
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10.5)
                    font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    font.color.rgb = RGBColor(0, 0, 0)
    for r in range(row_num_1):  # 表格中内容置为居中
        for c in range(5):
            table1.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    h3 = document.add_heading("", level=3)
    r3 = h3.add_run('1.' + str(strcnt) + '.3' + "调用函数")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.bold = False
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    r3.font.color.rgb = RGBColor(0, 0, 0)

    row_num_2 = len(funcs[cnt - 1].split(',')) + 1  #调用函数表格行数
    if row_num_2 == 2:
        if funcs[cnt - 1].split(',') == ['']:
            row_num_2 = 1
    table2 = document.add_table(rows=row_num_2, cols=3, style='Table Grid')  # 添加调用函数表格
    table2.autofit = True
    table2.cell(0, 0).text = "序号"
    table2.cell(0, 1).text = "调用函数名称"
    table2.cell(0, 2).text = "说明"
    for x in range(row_num_2 - 1):
        table2.cell(x+1, 0).text = str(x+1)   # 表格每行序号
    if row_num_2 > 1:
        x = 1
        for b in funcs[cnt - 1].split(','):
            table2.cell(x, 1).text = b
            x = x + 1

    for row in table2.rows:  # 设置表格中字体格式
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10.5)
                    font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    font.color.rgb = RGBColor(0, 0, 0)
    for r in range(row_num_2):  # 表格中内容置为居中
        for c in range(3):
            table2.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    h3 = document.add_heading("", level=3)
    r3 = h3.add_run('1.' + str(strcnt) + '.4' + "处理过程")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.bold = False
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    r3.font.color.rgb = RGBColor(0, 0, 0)

    p1 = document.add_paragraph()
    r_p1 = p1.add_run("  " + m['process'])
    r_p1.font.name = 'Times New Roman'
    r_p1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r_p1.font.size = Pt(12)
    r_p1.font.color.rgb = RGBColor(0, 0, 0)

    h3 = document.add_heading("", level=3)
    r3 = h3.add_run('1.' + str(strcnt) + '.5' + "函数输出")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.bold = False
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    r3.font.color.rgb = RGBColor(0, 0, 0)

    row_num_3 = 2
    table3 = document.add_table(rows=row_num_3, cols=5, style='Table Grid')  # 添加函数输出表格
    table3.autofit = True
    table3.cell(0, 0).text = "序号"
    table3.cell(0, 1).text = "数据名称"
    table3.cell(0, 2).text = "数据类型"
    table3.cell(0, 3).text = "单位"
    table3.cell(0, 4).text = "数据范围"
    table3.cell(1, 0).text = str(1)
    table3.cell(1, 1).text = "函数返回值"
    table3.cell(1, 2).text = m['return']

    for row in table3.rows:  # 设置表格中字体格式
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10.5)
                    font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    font.color.rgb = RGBColor(0, 0, 0)
    for r in range(row_num_3):  # 表格中内容置为居中
        for c in range(5):
            table3.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    cnt = cnt+1

filepath = savepath + '/'+outputname + '.docx'
document.save(filepath)
